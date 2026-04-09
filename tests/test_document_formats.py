"""Tests for multi-format document support (DOCX, DOC, RTF, ODT)."""

import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from pdf2txt import (
    SUPPORTED_EXTENSIONS,
    find_documents,
    extract_text_from_docx,
    extract_text,
    create_markdown,
    _heading_level,
    _table_to_markdown,
    _paragraph_has_page_break,
    check_libreoffice_available,
    process_document,
)


class TestSupportedExtensions:
    """Test the SUPPORTED_EXTENSIONS constant."""

    def test_contains_pdf(self):
        assert '.pdf' in SUPPORTED_EXTENSIONS

    def test_contains_docx(self):
        assert '.docx' in SUPPORTED_EXTENSIONS

    def test_contains_doc(self):
        assert '.doc' in SUPPORTED_EXTENSIONS

    def test_contains_rtf(self):
        assert '.rtf' in SUPPORTED_EXTENSIONS

    def test_contains_odt(self):
        assert '.odt' in SUPPORTED_EXTENSIONS


class TestFindDocuments:
    """Test find_documents() with multiple formats."""

    def test_finds_pdf_files(self, tmp_path):
        (tmp_path / "test.pdf").touch()
        (tmp_path / "test.txt").touch()
        result = find_documents(tmp_path, quiet=True)
        assert len(result) == 1
        assert result[0].suffix == '.pdf'

    def test_finds_docx_files(self, tmp_path):
        (tmp_path / "test.docx").touch()
        result = find_documents(tmp_path, quiet=True)
        assert len(result) == 1
        assert result[0].suffix == '.docx'

    def test_finds_mixed_formats(self, tmp_path):
        (tmp_path / "a.pdf").touch()
        (tmp_path / "b.docx").touch()
        (tmp_path / "c.doc").touch()
        (tmp_path / "d.txt").touch()
        result = find_documents(tmp_path, quiet=True)
        assert len(result) == 3

    def test_format_filter(self, tmp_path):
        (tmp_path / "a.pdf").touch()
        (tmp_path / "b.docx").touch()
        result = find_documents(tmp_path, quiet=True, formats={'.docx'})
        assert len(result) == 1
        assert result[0].suffix == '.docx'

    def test_recursive_finds_nested(self, tmp_path):
        sub = tmp_path / "sub"
        sub.mkdir()
        (sub / "nested.docx").touch()
        (tmp_path / "top.pdf").touch()
        result = find_documents(tmp_path, recursive=True, quiet=True)
        assert len(result) == 2

    def test_case_insensitive(self, tmp_path):
        (tmp_path / "test.DOCX").touch()
        result = find_documents(tmp_path, quiet=True)
        assert len(result) == 1

    def test_empty_directory(self, tmp_path):
        result = find_documents(tmp_path, quiet=True)
        assert len(result) == 0

    def test_shuffle(self, tmp_path):
        for i in range(20):
            (tmp_path / f"doc{i:02d}.pdf").touch()
        result = find_documents(tmp_path, quiet=True, shuffle=True)
        assert len(result) == 20


class TestHeadingLevel:
    """Test _heading_level() helper."""

    def test_heading_1(self):
        assert _heading_level("Heading 1") == 1

    def test_heading_2(self):
        assert _heading_level("Heading 2") == 2

    def test_heading_3(self):
        assert _heading_level("Heading 3") == 3

    def test_no_number_defaults_to_1(self):
        assert _heading_level("Heading") == 1


class TestExtractTextFromDocx:
    """Test DOCX extraction with real python-docx documents."""

    def _create_docx(self, tmp_path, paragraphs=None, filename="test.docx"):
        """Create a simple DOCX file for testing."""
        import docx
        doc = docx.Document()
        if paragraphs:
            for text in paragraphs:
                doc.add_paragraph(text)
        path = tmp_path / filename
        doc.save(str(path))
        return path

    def test_basic_extraction(self, tmp_path):
        path = self._create_docx(tmp_path, ["Hello world", "Second paragraph"])
        result = extract_text_from_docx(path)
        assert len(result) == 1  # Single section (no page breaks)
        assert "Hello world" in result[0]
        assert "Second paragraph" in result[0]

    def test_empty_docx(self, tmp_path):
        path = self._create_docx(tmp_path)
        result = extract_text_from_docx(path)
        assert len(result) == 1
        assert result[0] == ''

    def test_heading_conversion(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_heading("Main Title", level=1)
        doc.add_paragraph("Some content")
        doc.add_heading("Subtitle", level=2)
        path = tmp_path / "headings.docx"
        doc.save(str(path))

        result = extract_text_from_docx(path)
        assert "# Main Title" in result[0]
        assert "## Subtitle" in result[0]

    def test_table_extraction(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        doc.add_paragraph("After table")
        path = tmp_path / "table.docx"
        doc.save(str(path))

        result = extract_text_from_docx(path)
        assert "| A | B |" in result[0]
        assert "| C | D |" in result[0]
        assert "| --- | --- |" in result[0]

    def test_page_break_splits_sections(self, tmp_path):
        import docx
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = docx.Document()
        doc.add_paragraph("Page 1 content")

        # Add page break
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._element.append(br)
        run.add_text("Page 2 content")

        path = tmp_path / "pagebreak.docx"
        doc.save(str(path))

        result = extract_text_from_docx(path)
        assert len(result) == 2
        assert "Page 1 content" in result[0]
        assert "Page 2 content" in result[1]

    def test_stats_tracking(self, tmp_path):
        from pdf2txt import ProcessingStats
        stats = ProcessingStats()
        path = self._create_docx(tmp_path, ["Content"])
        extract_text_from_docx(path, stats=stats)
        assert stats.total_pages == 1
        assert stats.processed_pages == 1


class TestExtractTextDispatcher:
    """Test the extract_text() format dispatcher."""

    def test_dispatches_docx(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Test content")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        result = extract_text(path)
        assert "Test content" in result[0]

    def test_unsupported_format_raises(self, tmp_path):
        path = tmp_path / "test.xyz"
        path.touch()
        with pytest.raises(ValueError, match="Unsupported format"):
            extract_text(path)

    def test_doc_without_libreoffice(self, tmp_path):
        path = tmp_path / "test.doc"
        path.write_bytes(b"fake doc content")
        with patch('pdf2txt.check_libreoffice_available', return_value=False):
            with pytest.raises(RuntimeError, match="LibreOffice is required"):
                extract_text(path)


class TestCreateMarkdown:
    """Test create_markdown() with page_label parameter."""

    def test_default_page_label(self):
        path = Path("/tmp/test.pdf")
        result = create_markdown(path, ["Page 1", "Page 2"])
        assert "*Page 2*" in result

    def test_section_label(self):
        path = Path("/tmp/test.docx")
        result = create_markdown(path, ["Sec 1", "Sec 2"], page_label="Section")
        assert "*Section 2*" in result

    def test_single_section_no_separator(self):
        path = Path("/tmp/test.docx")
        result = create_markdown(path, ["Only content"], page_label="Section")
        assert "*Section" not in result
        assert "Only content" in result

    def test_source_header(self):
        path = Path("/tmp/mydoc.docx")
        result = create_markdown(path, ["Content"])
        assert "# mydoc" in result
        assert "> Source: /tmp/mydoc.docx" in result


class TestProcessDocument:
    """Test process_document() with DOCX files."""

    def test_creates_md_from_docx(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Hello from DOCX")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        success, message, _ = process_document(path, overwrite=False, dry_run=False)
        assert success
        assert "Created:" in message

        md_path = tmp_path / "test.md"
        assert md_path.exists()
        content = md_path.read_text()
        assert "Hello from DOCX" in content

    def test_skip_existing_md(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Content")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        (tmp_path / "test.md").write_text("existing")

        success, message, _ = process_document(path, overwrite=False, dry_run=False)
        assert not success
        assert "Skipped" in message

    def test_overwrite_existing_md(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("New content")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        (tmp_path / "test.md").write_text("old")

        success, message, _ = process_document(path, overwrite=True, dry_run=False)
        assert success
        assert "Created:" in message
        assert "New content" in (tmp_path / "test.md").read_text()

    def test_dry_run(self, tmp_path):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Content")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        success, message, _ = process_document(path, overwrite=False, dry_run=True)
        assert success
        assert "Would create" in message
        assert not (tmp_path / "test.md").exists()
